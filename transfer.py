import os
from docx import Document

def add_to_word(file_path, doc):
    """将文件内容添加到Word文档中"""
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
        doc.add_paragraph(content)
        doc.add_paragraph('')  # 添加空行作为文件间的分隔

def create_word_document(directory, output_path):
    """创建Word文档，包含目录中所有.py文件的内容"""
    # 创建一个新的Word文档
    doc = Document()
    doc.add_heading('Python Files Content', level=1)

    # 遍历目录下的所有文件
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.py'):
                file_path = os.path.join(root, file)
                add_to_word(file_path, doc)

    # 保存Word文档
    doc.save(output_path)
    print(f'Word文档已创建：{output_path}')

# 设置要扫描的目录和输出文档的路径
directory_to_scan = './'
output_document_path = 'python_files_content.docx'

# 创建Word文档
create_word_document(directory_to_scan, output_document_path)